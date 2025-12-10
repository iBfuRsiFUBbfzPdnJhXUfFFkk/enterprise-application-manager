import re
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from core.models.recommendation import Recommendation
from core.views.application.utilities.application_docx_helpers import (
    add_header_footer as _add_header_footer,
    add_toc_placeholder as _add_toc_placeholder,
    set_narrow_margins as _set_narrow_margins,
)


def set_narrow_margins(document: Document) -> None:
    """Reuse application helper for narrow margins."""
    _set_narrow_margins(document)


def add_header_footer(document: Document, title: str, person_name: str, date_str: str) -> None:
    """Add custom header with person (left), title (center), date (right) and footer with page numbers."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor, Inches
    from docx.table import _Cell

    section = document.sections[0]

    # Add header with a table for proper three-column layout
    header = section.header

    # Clear existing paragraphs
    for para in header.paragraphs:
        para.clear()

    # Create a 1-row, 3-column table in the header
    table = header.add_table(rows=1, cols=3, width=Inches(7.5))
    table.autofit = False

    # Set column widths
    table.columns[0].width = Inches(2.5)  # Left column (person)
    table.columns[1].width = Inches(2.5)  # Center column (title)
    table.columns[2].width = Inches(2.5)  # Right column (date)

    # Get cells
    cells = table.rows[0].cells

    # Left cell - Person name
    left_para = cells[0].paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_para.add_run(person_name)
    left_run.font.size = Pt(9)
    left_run.font.color.rgb = RGBColor(100, 100, 100)

    # Center cell - Title
    center_para = cells[1].paragraphs[0]
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_run = center_para.add_run(title)
    center_run.font.size = Pt(9)
    center_run.font.color.rgb = RGBColor(100, 100, 100)
    center_run.bold = True

    # Right cell - Date
    right_para = cells[2].paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_para.add_run(date_str)
    right_run.font.size = Pt(9)
    right_run.font.color.rgb = RGBColor(100, 100, 100)

    # Remove table borders
    from docx.oxml.shared import OxmlElement, qn
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Set borders to none
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Add footer with page numbers
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Add field for page number
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_toc_placeholder(document: Document) -> None:
    """Reuse application helper for TOC placeholder."""
    _add_toc_placeholder(document)


def add_markdown_content(document: Document, text: str, font_size: int = 9) -> None:
    """Add markdown-formatted text to the document with proper styling."""
    if not text or not text.strip():
        return

    lines = text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Handle headers
        if line.startswith('###'):
            para = document.add_paragraph()
            run = para.add_run(line[3:].strip())
            run.font.size = Pt(font_size + 2)
            run.bold = True
            i += 1
            continue
        elif line.startswith('##'):
            para = document.add_paragraph()
            run = para.add_run(line[2:].strip())
            run.font.size = Pt(font_size + 3)
            run.bold = True
            i += 1
            continue
        elif line.startswith('#'):
            para = document.add_paragraph()
            run = para.add_run(line[1:].strip())
            run.font.size = Pt(font_size + 4)
            run.bold = True
            i += 1
            continue

        # Handle unordered lists
        if line.strip().startswith(('- ', '* ', '+ ')):
            para = document.add_paragraph(style='List Bullet')
            add_inline_formatting(para, line.strip()[2:], font_size)
            i += 1
            continue

        # Handle ordered lists
        match = re.match(r'^(\d+)\.\s+(.*)$', line.strip())
        if match:
            para = document.add_paragraph(style='List Number')
            add_inline_formatting(para, match.group(2), font_size)
            i += 1
            continue

        # Handle code blocks
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                para = document.add_paragraph()
                run = para.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(font_size - 1)
                para_format = para.paragraph_format
                para_format.left_indent = Pt(20)
                # Add background shading
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F3F4F6')
                para._element.get_or_add_pPr().append(shading_elm)
            i += 1
            continue

        # Handle blockquotes
        if line.strip().startswith('>'):
            para = document.add_paragraph()
            add_inline_formatting(para, line.strip()[1:].strip(), font_size)
            para_format = para.paragraph_format
            para_format.left_indent = Pt(20)
            # Add left border for blockquote
            pPr = para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '10')
            left.set(qn('w:color'), 'D1D5DB')
            pBdr.append(left)
            pPr.append(pBdr)
            i += 1
            continue

        # Handle horizontal rules
        if line.strip() in ('---', '***', '___'):
            para = document.add_paragraph()
            pPr = para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'D1D5DB')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Handle regular paragraphs
        if line.strip():
            para = document.add_paragraph()
            add_inline_formatting(para, line, font_size)

        i += 1


def add_inline_formatting(paragraph, text: str, font_size: int = 9) -> None:
    """Add text with inline markdown formatting (bold, italic, code, links) to a paragraph."""
    # Pattern to match markdown inline elements
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)|'  # Bold + Italic
        r'(\*\*(.+?)\*\*)|'  # Bold
        r'(__(.+?)__)|'  # Bold (underscore)
        r'(\*(.+?)\*)|'  # Italic
        r'(_(.+?)_)|'  # Italic (underscore)
        r'(`(.+?)`)|'  # Inline code
        r'(\[(.+?)\]\((.+?)\))'  # Links
    )

    last_end = 0
    for match in pattern.finditer(text):
        # Add text before match
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.size = Pt(font_size)

        # Determine match type and add formatted run
        if match.group(1):  # Bold + Italic (***text***)
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
            run.font.size = Pt(font_size)
        elif match.group(3):  # Bold (**text**)
            run = paragraph.add_run(match.group(4))
            run.bold = True
            run.font.size = Pt(font_size)
        elif match.group(5):  # Bold (__text__)
            run = paragraph.add_run(match.group(6))
            run.bold = True
            run.font.size = Pt(font_size)
        elif match.group(7):  # Italic (*text*)
            run = paragraph.add_run(match.group(8))
            run.italic = True
            run.font.size = Pt(font_size)
        elif match.group(9):  # Italic (_text_)
            run = paragraph.add_run(match.group(10))
            run.italic = True
            run.font.size = Pt(font_size)
        elif match.group(11):  # Inline code (`code`)
            run = paragraph.add_run(match.group(12))
            run.font.name = 'Courier New'
            run.font.size = Pt(font_size - 1)
            # Add background to inline code
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F3F4F6')
            run._element.get_or_add_rPr().append(shading_elm)
        elif match.group(13):  # Link [text](url)
            run = paragraph.add_run(match.group(14))
            run.font.color.rgb = RGBColor(59, 130, 246)
            run.underline = True
            run.font.size = Pt(font_size)
            # Note: actual hyperlink functionality would require more complex DOCX manipulation

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.size = Pt(font_size)


def add_recommendation_intro(document: Document) -> None:
    """Add introductory paragraph about recommendations."""
    from docx.shared import Pt

    intro_para = document.add_paragraph()
    intro_run = intro_para.add_run(
        "This document contains a technical recommendation for improving systems, processes, or practices. "
        "Recommendations are tracked to ensure continuous improvement and documentation of best practices. "
        "Each recommendation includes detailed analysis of benefits, risks, and implementation considerations."
    )
    intro_run.font.size = Pt(10)
    intro_run.italic = True

    # Add spacing after intro
    document.add_paragraph()


def add_title_page(document: Document, username: str, total_count: int) -> None:
    """Add title page with metadata."""
    title = document.add_heading("Recommendations Export", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    table = document.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"

    table.rows[0].cells[0].text = "Generated By"
    table.rows[0].cells[1].text = username

    table.rows[1].cells[0].text = "Generated Date"
    now = datetime.now(timezone.utc)
    table.rows[1].cells[1].text = now.strftime("%Y-%m-%d %I:%M %p UTC")

    table.rows[2].cells[0].text = "Total Recommendations"
    table.rows[2].cells[1].text = str(total_count)

    # Format table
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    document.add_page_break()


def add_recommendation_section(document: Document, recommendation: Recommendation) -> None:
    """Add a formatted recommendation section to the document."""
    from core.utilities.get_name_acronym import get_name_acronym

    # Metadata table - only show related entities
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"

    # Related entities
    if recommendation.application:
        row = table.add_row()
        row.cells[0].text = "Related Application"
        row.cells[1].text = get_name_acronym(
            acronym=recommendation.application.acronym, name=recommendation.application.name
        )

    if recommendation.project:
        row = table.add_row()
        row.cells[0].text = "Related Project"
        row.cells[1].text = recommendation.project.name if hasattr(recommendation.project, "name") else str(recommendation.project)

    if recommendation.estimation:
        row = table.add_row()
        row.cells[0].text = "Related Estimation"
        row.cells[1].text = recommendation.estimation.name if hasattr(recommendation.estimation, "name") else str(recommendation.estimation)

    # Format table cells
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Add spacing
    document.add_paragraph()

    # Description section
    if recommendation.description:
        document.add_heading("Description", level=2)
        add_markdown_content(document, recommendation.description)

    # Rationale section
    if recommendation.rationale:
        document.add_heading("Rationale", level=2)
        add_markdown_content(document, recommendation.rationale)

    # Expected Benefits section
    if recommendation.benefits:
        document.add_heading("Expected Benefits", level=2)
        add_markdown_content(document, recommendation.benefits)

    # Potential Risks section
    if recommendation.risks:
        document.add_heading("Potential Risks", level=2)
        add_markdown_content(document, recommendation.risks)

    # Additional Notes section
    if recommendation.comment:
        document.add_heading("Notes", level=2)
        add_markdown_content(document, recommendation.comment)
