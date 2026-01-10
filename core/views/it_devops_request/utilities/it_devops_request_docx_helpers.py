import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from core.models.it_devops_request import ITDevOpsRequest
from core.utilities.pdf_helpers import (
    MAX_EMBED_SIZE,
    convert_pdf_to_images,
    format_file_size,
    get_attachment_handler,
)
from core.views.recommendation.utilities.recommendation_docx_helpers import (
    add_header_footer as _add_header_footer,
    add_hyperlink,
    add_markdown_content,
    set_narrow_margins as _set_narrow_margins,
)


def set_narrow_margins(document: Document) -> None:
    """Reuse recommendation helper for narrow margins."""
    _set_narrow_margins(document)


def add_header_footer(document: Document, title: str, person_name: str, date_str: str) -> None:
    """Reuse recommendation helper for header/footer with custom title."""
    _add_header_footer(document, title, person_name, date_str)


def add_request_intro(document: Document) -> None:
    """Add introductory paragraph about IT/DevOps requests."""
    intro_para = document.add_paragraph()
    intro_run = intro_para.add_run(
        "This document contains an IT/DevOps request for system changes, infrastructure modifications, "
        "or technical support. Each request includes detailed specifications, justification, expected outcomes, "
        "and a complete timeline of updates and progress."
    )
    intro_run.font.size = Pt(10)
    intro_run.italic = True


def add_links_section(document: Document, request: ITDevOpsRequest) -> None:
    """Add links section to the document with clickable hyperlinks."""
    if not request.links.exists():
        return

    document.add_heading("Related Links", level=2)

    for link in request.links.all():
        para = document.add_paragraph(style='List Bullet')

        # Add link name in bold
        name_run = para.add_run(link.name)
        name_run.bold = True
        name_run.font.size = Pt(9)

        # Add line break
        para.add_run('\n')

        # Add full URL as clickable hyperlink (using original url, NOT short_code)
        add_hyperlink(para, link.url, link.url, font_size=9)

        # Add comment if exists
        if link.comment:
            para.add_run('\n')
            comment_run = para.add_run(link.comment)
            comment_run.font.size = Pt(8)
            comment_run.italic = True


def add_request_section(document: Document, request: ITDevOpsRequest) -> None:
    """Add a formatted IT/DevOps request section to the document."""
    from core.utilities.get_name_acronym import get_name_acronym

    # Document ID - prominently displayed at the top
    document.add_heading(f"Document ID: {request.document_id}", level=1)

    # Request name as subtitle
    if request.name:
        name_para = document.add_paragraph()
        name_run = name_para.add_run(request.name)
        name_run.font.size = Pt(14)
        name_run.bold = True
        name_para.paragraph_format.space_after = Pt(12)

    # Overview table - status, priority, reference number, dates
    document.add_heading("Request Overview", level=2)
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"

    # Status
    row = table.add_row()
    row.cells[0].text = "Status"
    row.cells[1].text = request.get_status_display() if hasattr(request, "get_status_display") else request.status

    # Priority
    row = table.add_row()
    row.cells[0].text = "Priority"
    row.cells[1].text = request.get_priority_display() if hasattr(request, "get_priority_display") else request.priority

    # Reference number
    if request.reference_number:
        row = table.add_row()
        row.cells[0].text = "Reference/Ticket Number"
        row.cells[1].text = request.reference_number

    # Date requested
    if request.date_requested:
        row = table.add_row()
        row.cells[0].text = "Date Requested"
        row.cells[1].text = request.date_requested.strftime("%Y-%m-%d")

    # Date due
    if request.date_due:
        row = table.add_row()
        row.cells[0].text = "Date Due"
        row.cells[1].text = request.date_due.strftime("%Y-%m-%d")

    # Date completed
    if request.date_completed:
        row = table.add_row()
        row.cells[0].text = "Date Completed"
        row.cells[1].text = request.date_completed.strftime("%Y-%m-%d")

    # Format table cells
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # People section
    document.add_heading("People", level=2)
    people_table = document.add_table(rows=0, cols=2)
    people_table.style = "Light Grid Accent 1"

    # Requester
    if request.person_requester:
        row = people_table.add_row()
        row.cells[0].text = "Requester"
        row.cells[1].text = str(request.person_requester)

    # Assignee
    if request.person_assignee:
        row = people_table.add_row()
        row.cells[0].text = "Assignee"
        row.cells[1].text = str(request.person_assignee)

    # Approver
    if request.person_approver:
        row = people_table.add_row()
        row.cells[0].text = "Approver"
        row.cells[1].text = str(request.person_approver)

    # Format people table
    for row in people_table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Related entities
    if request.application or request.project:
        document.add_heading("Related Entities", level=2)
        entities_table = document.add_table(rows=0, cols=2)
        entities_table.style = "Light Grid Accent 1"

        if request.application:
            row = entities_table.add_row()
            row.cells[0].text = "Application"
            row.cells[1].text = get_name_acronym(
                acronym=request.application.acronym, name=request.application.name
            )

        if request.project:
            row = entities_table.add_row()
            row.cells[0].text = "Project"
            row.cells[1].text = request.project.name if hasattr(request.project, "name") else str(request.project)

        # Format entities table
        for row in entities_table.rows:
            row.cells[0].paragraphs[0].runs[0].bold = True
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # Description section
    if request.description:
        document.add_heading("Description", level=2)
        add_markdown_content(document, request.description)

    # Justification section
    if request.justification:
        document.add_heading("Justification", level=2)
        add_markdown_content(document, request.justification)

    # Expected outcome section
    if request.expected_outcome:
        document.add_heading("Expected Outcome", level=2)
        add_markdown_content(document, request.expected_outcome)

    # Additional notes section
    if request.comment:
        document.add_heading("Additional Notes", level=2)
        add_markdown_content(document, request.comment)


def add_updates_section(document: Document, updates) -> None:
    """Add updates timeline section to the document."""
    if not updates:
        return

    document.add_heading("Updates Timeline", level=2)

    for update in updates:
        # Create a paragraph for each update
        para = document.add_paragraph()

        # Timestamp and author in bold
        timestamp_str = update.datetime_created.strftime("%Y-%m-%d %H:%M:%S UTC") if update.datetime_created else "Unknown date"
        author_str = str(update.person_author) if update.person_author else "Unknown author"
        header_text = f"{timestamp_str} - {author_str}"

        if update.is_internal_note:
            header_text += " (Internal Note)"

        header_run = para.add_run(header_text)
        header_run.bold = True
        header_run.font.size = Pt(9)

        # Comment text
        if update.comment:
            para.add_run("\n")
            comment_run = para.add_run(update.comment)
            comment_run.font.size = Pt(9)

        # Add spacing
        para.paragraph_format.space_after = Pt(8)

        # Add horizontal line after each update except the last
        if update != updates[len(updates) - 1]:
            hr_para = document.add_paragraph()
            hr_para.paragraph_format.space_after = Pt(4)
            hr_para.paragraph_format.space_before = Pt(4)
            pPr = hr_para._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "D1D5DB")
            pBdr.append(bottom)
            pPr.append(pBdr)


def add_attachments_section(document: Document, request: ITDevOpsRequest) -> None:
    """Add attachments section with embedded content (text, images, PDFs as images)."""
    import mimetypes
    from PIL import Image

    if not request.attachments.exists():
        return

    document.add_heading("Attachments", level=2)

    first_attachment = True
    for attachment in request.attachments.all():
        # Skip if no file attached
        if not attachment.file:
            continue

        # Add page break before each attachment except the first
        if not first_attachment:
            document.add_page_break()
        first_attachment = False

        # Get file info from MinIO
        filename = attachment.get_filename()
        file_size = attachment.get_file_size()

        # Guess content type from filename
        content_type, _ = mimetypes.guess_type(filename) if filename else (None, None)
        if not content_type:
            content_type = 'application/octet-stream'

        # Add filename as heading
        document.add_heading(filename or "Untitled", level=3)

        # Check file size
        if file_size and file_size > MAX_EMBED_SIZE:
            para = document.add_paragraph()
            para.add_run(f"File Type: {content_type}\n").bold = True
            para.add_run(f"File Size: {format_file_size(file_size)}\n\n")
            para.add_run("This file is too large to embed (>10MB).").italic = True
            continue

        handler = get_attachment_handler(content_type)

        if handler == "text":
            try:
                # Read file from MinIO
                attachment.file.open('rb')
                file_data = attachment.file.read()
                attachment.file.close()

                text = file_data.decode("utf-8", errors="ignore")
                # Add text content with monospace font
                lines = text.split("\n")
                for line in lines:
                    para = document.add_paragraph()
                    run = para.add_run(line if line.strip() else " ")
                    run.font.name = "Courier New"
                    run.font.size = Pt(8)
            except Exception:
                para = document.add_paragraph()
                para.add_run("Error: Unable to decode text content.").italic = True

        elif handler == "image":
            try:
                # Read file from MinIO
                attachment.file.open('rb')
                file_data = attachment.file.read()
                attachment.file.close()

                img = Image.open(BytesIO(file_data))
                # Resize if needed
                max_width = 6.5
                max_height = 8.0
                width_ratio = max_width / (img.width / 96)  # Convert pixels to inches (96 DPI)
                height_ratio = max_height / (img.height / 96)
                ratio = min(width_ratio, height_ratio, 1.0)  # Don't upscale

                img_stream = BytesIO()
                img.save(img_stream, format="PNG")
                img_stream.seek(0)

                document.add_picture(img_stream, width=Inches(img.width / 96 * ratio))
            except Exception:
                para = document.add_paragraph()
                para.add_run("Error: Unable to load image.").italic = True

        elif handler == "pdf":
            try:
                # Read file from MinIO
                attachment.file.open('rb')
                file_data = attachment.file.read()
                attachment.file.close()

                images = convert_pdf_to_images(file_data)
                if images:
                    for i, img in enumerate(images):
                        if i > 0:
                            para = document.add_paragraph()
                            para.add_run(f"\nPage {i + 1}").bold = True

                        # Resize to fit page
                        max_width = 6.5
                        max_height = 8.0
                        width_ratio = max_width / (img.width / 96)
                        height_ratio = max_height / (img.height / 96)
                        ratio = min(width_ratio, height_ratio, 1.0)

                        img_stream = BytesIO()
                        img.save(img_stream, format="PNG")
                        img_stream.seek(0)

                        document.add_picture(img_stream, width=Inches(img.width / 96 * ratio))
                else:
                    para = document.add_paragraph()
                    para.add_run("Error: Unable to convert PDF to images.").italic = True
            except Exception:
                para = document.add_paragraph()
                para.add_run("Error: Unable to process PDF.").italic = True

        else:
            # Unsupported file type - show metadata
            para = document.add_paragraph()
            para.add_run(f"File Type: {content_type}\n").bold = True
            if file_size:
                para.add_run(f"File Size: {format_file_size(file_size)}\n\n")
            para.add_run("This file type cannot be embedded in the document.").italic = True
