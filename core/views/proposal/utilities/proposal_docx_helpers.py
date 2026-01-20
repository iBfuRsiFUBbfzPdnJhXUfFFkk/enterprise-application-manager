from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from core.models.proposal import Proposal
from core.utilities.pdf_helpers import (
    MAX_EMBED_SIZE,
    convert_pdf_to_images,
    format_file_size,
    get_attachment_handler,
)
from core.views.recommendation.utilities.recommendation_docx_helpers import (
    add_header_footer as _add_header_footer,
    add_hyperlink,
    add_markdown_content,
    set_narrow_margins as _set_narrow_margins,
)
from docx.shared import Inches


def set_narrow_margins(document: Document) -> None:
    """Reuse recommendation helper for narrow margins."""
    _set_narrow_margins(document)


def add_header_footer(document: Document, title: str, person_name: str, date_str: str) -> None:
    """Reuse recommendation helper for header/footer with custom title."""
    _add_header_footer(document, title, person_name, date_str)


def add_proposal_intro(document: Document) -> None:
    """Add introductory paragraph about proposals."""
    intro_para = document.add_paragraph()
    intro_run = intro_para.add_run(
        "This document contains a formal proposal outlining a proposed initiative, project, or change. "
        "It includes an executive summary, problem statement, proposed solution, benefits analysis, "
        "risk assessment, and implementation timeline."
    )
    intro_run.font.size = Pt(10)
    intro_run.italic = True


def add_links_section(document: Document, proposal: Proposal) -> None:
    """Add links section to the document with clickable hyperlinks."""
    if not proposal.links.exists():
        return

    document.add_heading("Related Links", level=2)

    for link in proposal.links.all():
        para = document.add_paragraph(style='List Bullet')

        name_run = para.add_run(link.name)
        name_run.bold = True
        name_run.font.size = Pt(9)

        para.add_run('\n')
        add_hyperlink(para, link.url, link.url, font_size=9)

        if link.comment:
            para.add_run('\n')
            comment_run = para.add_run(link.comment)
            comment_run.font.size = Pt(8)
            comment_run.italic = True


def add_proposal_section(document: Document, proposal: Proposal) -> None:
    """Add a formatted proposal section to the document."""
    from core.utilities.get_name_acronym import get_name_acronym

    # Document ID
    document.add_heading(f"Document ID: {proposal.document_id}", level=1)

    # Proposal name as subtitle
    if proposal.name:
        name_para = document.add_paragraph()
        name_run = name_para.add_run(proposal.name)
        name_run.font.size = Pt(14)
        name_run.bold = True
        name_para.paragraph_format.space_after = Pt(12)

    # Overview table
    document.add_heading("Proposal Overview", level=2)
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"

    # Status
    if proposal.status:
        row = table.add_row()
        row.cells[0].text = "Status"
        row.cells[1].text = proposal.get_status_display() if hasattr(proposal, "get_status_display") else proposal.status

    # Priority
    if proposal.priority:
        row = table.add_row()
        row.cells[0].text = "Priority"
        row.cells[1].text = proposal.get_priority_display() if hasattr(proposal, "get_priority_display") else proposal.priority

    # Version
    if proposal.version:
        row = table.add_row()
        row.cells[0].text = "Version"
        row.cells[1].text = proposal.version

    # Reference number
    if proposal.reference_number:
        row = table.add_row()
        row.cells[0].text = "Reference Number"
        row.cells[1].text = proposal.reference_number

    # Format table cells
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # People section
    if proposal.person_author or proposal.person_reviewer or proposal.person_approver:
        document.add_heading("People", level=2)
        people_table = document.add_table(rows=0, cols=2)
        people_table.style = "Light Grid Accent 1"

        if proposal.person_author:
            row = people_table.add_row()
            row.cells[0].text = "Author"
            row.cells[1].text = str(proposal.person_author)

        if proposal.person_reviewer:
            row = people_table.add_row()
            row.cells[0].text = "Reviewer"
            row.cells[1].text = str(proposal.person_reviewer)

        if proposal.person_approver:
            row = people_table.add_row()
            row.cells[0].text = "Approver"
            row.cells[1].text = str(proposal.person_approver)

        for row in people_table.rows:
            row.cells[0].paragraphs[0].runs[0].bold = True
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # Key dates
    has_dates = any([
        proposal.date_created,
        proposal.date_submitted,
        proposal.date_review_completed,
        proposal.date_decision,
        proposal.date_implementation_target,
    ])

    if has_dates:
        document.add_heading("Key Dates", level=2)
        dates_table = document.add_table(rows=0, cols=2)
        dates_table.style = "Light Grid Accent 1"

        if proposal.date_created:
            row = dates_table.add_row()
            row.cells[0].text = "Created"
            row.cells[1].text = proposal.date_created.strftime("%Y-%m-%d")

        if proposal.date_submitted:
            row = dates_table.add_row()
            row.cells[0].text = "Submitted"
            row.cells[1].text = proposal.date_submitted.strftime("%Y-%m-%d")

        if proposal.date_review_completed:
            row = dates_table.add_row()
            row.cells[0].text = "Review Completed"
            row.cells[1].text = proposal.date_review_completed.strftime("%Y-%m-%d")

        if proposal.date_decision:
            row = dates_table.add_row()
            row.cells[0].text = "Decision Date"
            row.cells[1].text = proposal.date_decision.strftime("%Y-%m-%d")

        if proposal.date_implementation_target:
            row = dates_table.add_row()
            row.cells[0].text = "Target Implementation"
            row.cells[1].text = proposal.date_implementation_target.strftime("%Y-%m-%d")

        for row in dates_table.rows:
            row.cells[0].paragraphs[0].runs[0].bold = True
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # Related entities
    if proposal.application or proposal.project:
        document.add_heading("Related Entities", level=2)
        entities_table = document.add_table(rows=0, cols=2)
        entities_table.style = "Light Grid Accent 1"

        if proposal.application:
            row = entities_table.add_row()
            row.cells[0].text = "Application"
            row.cells[1].text = get_name_acronym(
                acronym=proposal.application.acronym, name=proposal.application.name
            )

        if proposal.project:
            row = entities_table.add_row()
            row.cells[0].text = "Project"
            row.cells[1].text = proposal.project.name if hasattr(proposal.project, "name") else str(proposal.project)

        for row in entities_table.rows:
            row.cells[0].paragraphs[0].runs[0].bold = True
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # Content sections
    if proposal.executive_summary:
        document.add_heading("Executive Summary", level=2)
        add_markdown_content(document, proposal.executive_summary)

    if proposal.problem_statement:
        document.add_heading("Problem Statement", level=2)
        add_markdown_content(document, proposal.problem_statement)

    if proposal.proposed_solution:
        document.add_heading("Proposed Solution", level=2)
        add_markdown_content(document, proposal.proposed_solution)

    if proposal.benefits:
        document.add_heading("Benefits", level=2)
        add_markdown_content(document, proposal.benefits)

    if proposal.risks_and_mitigations:
        document.add_heading("Risks & Mitigations", level=2)
        add_markdown_content(document, proposal.risks_and_mitigations)

    if proposal.timeline:
        document.add_heading("Timeline", level=2)
        add_markdown_content(document, proposal.timeline)

    if proposal.resources_required:
        document.add_heading("Resources Required", level=2)
        add_markdown_content(document, proposal.resources_required)

    if proposal.success_criteria:
        document.add_heading("Success Criteria", level=2)
        add_markdown_content(document, proposal.success_criteria)

    if proposal.comment:
        document.add_heading("Additional Notes", level=2)
        add_markdown_content(document, proposal.comment)


def add_updates_section(document: Document, updates) -> None:
    """Add updates timeline section to the document."""
    if not updates:
        return

    document.add_heading("Updates Timeline", level=2)

    for update in updates:
        para = document.add_paragraph()

        timestamp_str = update.datetime_created.strftime("%Y-%m-%d %H:%M:%S UTC") if update.datetime_created else "Unknown date"
        author_str = str(update.person_author) if update.person_author else "Unknown author"
        header_text = f"{timestamp_str} - {author_str}"

        if update.is_internal_note:
            header_text += " (Internal Note)"

        header_run = para.add_run(header_text)
        header_run.bold = True
        header_run.font.size = Pt(9)

        if update.comment:
            para.add_run("\n")
            comment_run = para.add_run(update.comment)
            comment_run.font.size = Pt(9)

        para.paragraph_format.space_after = Pt(8)

        if update != updates[len(updates) - 1]:
            hr_para = document.add_paragraph()
            hr_para.paragraph_format.space_after = Pt(4)
            hr_para.paragraph_format.space_before = Pt(4)
            pPr = hr_para._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "D1D5DB")
            pBdr.append(bottom)
            pPr.append(pBdr)


def add_attachments_section(document: Document, proposal: Proposal) -> None:
    """Add attachments section with embedded content."""
    import mimetypes
    from PIL import Image

    if not proposal.attachments.exists():
        return

    document.add_heading("Attachments", level=2)

    first_attachment = True
    for attachment in proposal.attachments.all():
        if not attachment.file:
            continue

        if not first_attachment:
            document.add_page_break()
        first_attachment = False

        filename = attachment.get_filename()
        file_size = attachment.get_file_size()

        content_type, _ = mimetypes.guess_type(filename) if filename else (None, None)
        if not content_type:
            content_type = 'application/octet-stream'

        document.add_heading(filename or "Untitled", level=3)

        if file_size and file_size > MAX_EMBED_SIZE:
            para = document.add_paragraph()
            para.add_run(f"File Type: {content_type}\n").bold = True
            para.add_run(f"File Size: {format_file_size(file_size)}\n\n")
            para.add_run("This file is too large to embed (>10MB).").italic = True
            continue

        handler = get_attachment_handler(content_type)

        if handler == "text":
            try:
                attachment.file.open('rb')
                file_data = attachment.file.read()
                attachment.file.close()

                text = file_data.decode("utf-8", errors="ignore")
                lines = text.split("\n")
                for line in lines:
                    para = document.add_paragraph()
                    run = para.add_run(line if line.strip() else " ")
                    run.font.name = "Courier New"
                    run.font.size = Pt(8)
            except Exception:
                para = document.add_paragraph()
                para.add_run("Error: Unable to decode text content.").italic = True

        elif handler == "image":
            try:
                attachment.file.open('rb')
                file_data = attachment.file.read()
                attachment.file.close()

                img = Image.open(BytesIO(file_data))
                max_width = 6.5
                max_height = 8.0
                width_ratio = max_width / (img.width / 96)
                height_ratio = max_height / (img.height / 96)
                ratio = min(width_ratio, height_ratio, 1.0)

                img_stream = BytesIO()
                img.save(img_stream, format="PNG")
                img_stream.seek(0)

                document.add_picture(img_stream, width=Inches(img.width / 96 * ratio))
            except Exception:
                para = document.add_paragraph()
                para.add_run("Error: Unable to load image.").italic = True

        elif handler == "pdf":
            try:
                attachment.file.open('rb')
                file_data = attachment.file.read()
                attachment.file.close()

                images = convert_pdf_to_images(file_data)
                if images:
                    for i, img in enumerate(images):
                        if i > 0:
                            para = document.add_paragraph()
                            para.add_run(f"\nPage {i + 1}").bold = True

                        max_width = 6.5
                        max_height = 8.0
                        width_ratio = max_width / (img.width / 96)
                        height_ratio = max_height / (img.height / 96)
                        ratio = min(width_ratio, height_ratio, 1.0)

                        img_stream = BytesIO()
                        img.save(img_stream, format="PNG")
                        img_stream.seek(0)

                        document.add_picture(img_stream, width=Inches(img.width / 96 * ratio))
                else:
                    para = document.add_paragraph()
                    para.add_run("Error: Unable to convert PDF to images.").italic = True
            except Exception:
                para = document.add_paragraph()
                para.add_run("Error: Unable to process PDF.").italic = True

        else:
            para = document.add_paragraph()
            para.add_run(f"File Type: {content_type}\n").bold = True
            if file_size:
                para.add_run(f"File Size: {format_file_size(file_size)}\n\n")
            para.add_run("This file type cannot be embedded in the document.").italic = True
