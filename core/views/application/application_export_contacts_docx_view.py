from datetime import datetime, timezone
from io import BytesIO
from urllib.parse import urlparse

from django.contrib.auth.decorators import login_required
from django.http import HttpRequest, HttpResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from core.models.application import Application
from core.models.common.enums.lifecycle_choices import (
    LIFECYCLE_ACTIVE,
    LIFECYCLE_HYPER_CARE,
    LIFECYCLE_IN_DEPRECATION,
    LIFECYCLE_LIMITED_SUPPORT,
)
from core.models.person import Person
from core.utilities.get_name_acronym import get_name_acronym
from core.views.application.utilities.application_docx_helpers import (
    add_header_footer,
    set_narrow_margins,
)
from core.views.generic.generic_500 import generic_500


def add_person_contact_to_cell(cell, person: Person | None, document: Document, font_size: int = 8) -> None:
    """Add formatted person contact to cell with mailto: hyperlink for email."""
    if not person:
        cell.text = "—"
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
        return

    first = person.coerced_name_first or ""
    last = person.coerced_name_last or ""
    email = person.coerced_communication_email or ""

    if not first and not last:
        cell.text = "—"
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
        return

    # Get job title
    title = ""
    if person.job_title:
        title = person.job_title.name

    # Build name with title
    name = f"{last}, {first}".strip(", ")
    if title:
        name_line = f"{name} ({title})"
    else:
        name_line = name

    # Clear existing content
    cell.text = ""
    paragraph = cell.paragraphs[0]

    # Add name and title as regular text
    run = paragraph.add_run(name_line)
    run.font.size = Pt(font_size)

    # Add line break
    paragraph.add_run("\n")

    # Add email as mailto: hyperlink or plain text
    if email:
        # Create mailto: hyperlink
        mailto_url = f"mailto:{email}"

        # Add relationship for mailto hyperlink
        part = paragraph.part
        r_id = part.relate_to(mailto_url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create run element
        run_element = OxmlElement('w:r')

        # Create run properties
        rPr = OxmlElement('w:rPr')

        # Add hyperlink style
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)

        # Style as link (blue, underlined)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')  # Word hyperlink blue
        rPr.append(color)

        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)

        # Set font size
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size * 2))  # Half-points
        rPr.append(sz)

        run_element.append(rPr)

        # Add email text
        text_element = OxmlElement('w:t')
        text_element.text = email
        run_element.append(text_element)

        hyperlink.append(run_element)

        # Add hyperlink to paragraph
        paragraph._p.append(hyperlink)
    else:
        # No email - add plain text
        run = paragraph.add_run("(No email)")
        run.font.size = Pt(font_size)
        run.italic = True


def extract_hostname(url: str | None) -> str:
    """Extract hostname from URL for display purposes."""
    if not url:
        return "—"

    try:
        parsed = urlparse(url)
        hostname = parsed.netloc or parsed.path
        return hostname if hostname else "—"
    except Exception:
        return url


def add_hyperlink_to_cell(cell, url: str | None, document: Document, font_size: int = 8) -> None:
    """Add clickable hyperlink to table cell showing only hostname."""
    if not url or url == "—":
        cell.text = "—"
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
        return

    # Clear existing content
    cell.text = ""

    # Get or create paragraph
    paragraph = cell.paragraphs[0]

    # Add relationship for external hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create run element
    run_element = OxmlElement('w:r')

    # Create run properties
    rPr = OxmlElement('w:rPr')

    # Add hyperlink style
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # Style as link (blue, underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')  # Word hyperlink blue
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    # Set font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))  # Half-points
    rPr.append(sz)

    run_element.append(rPr)

    # Add text (hostname only)
    text_element = OxmlElement('w:t')
    text_element.text = extract_hostname(url)
    run_element.append(text_element)

    hyperlink.append(run_element)

    # Add hyperlink to paragraph
    paragraph._p.append(hyperlink)


def add_contact_list_title_page(document: Document, username: str, count: int) -> None:
    """Add title page with metadata for contact list export."""
    # Title
    title = document.add_heading('Application Notification Contact List', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    table = document.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'

    table.rows[0].cells[0].text = 'Generated By'
    table.rows[0].cells[1].text = username

    table.rows[1].cells[0].text = 'Generated Date'
    now = datetime.now(timezone.utc)
    table.rows[1].cells[1].text = now.strftime('%Y-%m-%d %I:%M %p UTC')

    table.rows[2].cells[0].text = 'Total Applications'
    table.rows[2].cells[1].text = str(count)

    # Format table
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Add disclaimer
    document.add_paragraph()  # Spacing
    disclaimer = document.add_paragraph()
    disclaimer_text = (
        'Note: This document is subject to change as projects and personnel assignments evolve. '
        f'The information contained herein represents the best available data as of {now.strftime("%Y-%m-%d")}.'
    )
    run = disclaimer.add_run(disclaimer_text)
    run.font.size = Pt(9)
    run.italic = True
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()


@login_required
def application_export_contacts_docx_view(request: HttpRequest) -> HttpResponse:
    """Export application contact list to DOCX for notification purposes."""
    try:
        # Define allowed lifecycle values for contact list
        allowed_lifecycle_values = [
            LIFECYCLE_ACTIVE,
            LIFECYCLE_HYPER_CARE,
            LIFECYCLE_LIMITED_SUPPORT,
            LIFECYCLE_IN_DEPRECATION,
        ]

        # Query applications with lifecycle filter and exclude EAM
        applications = Application.objects.filter(
            type_lifecycle__in=allowed_lifecycle_values
        ).exclude(
            acronym='EAM'
        ).select_related(
            'person_lead_developer',
            'person_lead_developer__job_title',
            'person_project_manager',
            'person_project_manager__job_title',
        ).order_by('-is_externally_facing', 'name', 'acronym')

        # Create document
        document = Document()
        set_narrow_margins(document)
        add_header_footer(document, 'Application Notification Contact List')

        # Add title page
        username = request.user.get_full_name() or request.user.username
        add_contact_list_title_page(document, username, applications.count())

        # Create main table
        num_apps = applications.count()
        table = document.add_table(rows=num_apps + 1, cols=8)
        table.style = 'Light Grid Accent 1'

        # Disable auto-fit to respect manual column widths
        table.autofit = False
        table.allow_autofit = False

        # Set static column widths (in inches, total 7.47" for portrait with 0.5" margins)
        column_widths = [
            1.05,  # Application
            0.4,   # External? (icon)
            1.35,  # Lead Developer
            1.35,  # Project Manager
            0.83,  # Dev
            0.83,  # Stage
            0.83,  # Prod
            0.83,  # External
        ]

        # Set column widths and also set width on every cell to enforce it
        for row in table.rows:
            for idx, width in enumerate(column_widths):
                row.cells[idx].width = Inches(width)

        # Build header row
        headers = ['Application', 'Ext', 'Lead Developer', 'Project Manager',
                   'Dev', 'Stage', 'Prod', 'External']
        header_row = table.rows[0]

        for col_idx, header_text in enumerate(headers):
            cell = header_row.cells[col_idx]
            cell.text = header_text

            # Format header cell
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

            # Set header background color (light blue)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'C8DCF0')  # RGB(200, 220, 240)
            cell._element.get_or_add_tcPr().append(shading_elm)

        # Build data rows
        for row_idx, app in enumerate(applications, start=1):
            row = table.rows[row_idx]

            # Column 0: Application name/acronym
            row.cells[0].text = get_name_acronym(acronym=app.acronym, name=app.name)

            # Column 1: External? (icon: ✓ for yes, blank for no)
            row.cells[1].text = "✓" if app.is_externally_facing else ""
            # Center the icon
            for paragraph in row.cells[1].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Column 2: Lead Developer with title and mailto: email
            add_person_contact_to_cell(row.cells[2], app.person_lead_developer, document)

            # Column 3: Project Manager with title and mailto: email
            add_person_contact_to_cell(row.cells[3], app.person_project_manager, document)

            # Column 4: Dev URL (hyperlink with hostname)
            add_hyperlink_to_cell(row.cells[4], app.link_development_server, document)

            # Column 5: Stage URL (hyperlink with hostname)
            add_hyperlink_to_cell(row.cells[5], app.link_staging_server, document)

            # Column 6: Prod URL (hyperlink with hostname)
            add_hyperlink_to_cell(row.cells[6], app.link_production_server, document)

            # Column 7: External Prod URL (hyperlink with hostname)
            add_hyperlink_to_cell(row.cells[7], app.link_production_server_external, document)

            # Format data cells (column 0 only - others already formatted)
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

            # Format External icon column (column 1)
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Slightly larger for visibility

            # Apply alternating row colors
            if row_idx % 2 == 0:
                # Even rows: light gray
                for cell in row.cells:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F5F5F5')  # RGB(245, 245, 245)
                    cell._element.get_or_add_tcPr().append(shading_elm)

        # Generate filename
        now = datetime.now(timezone.utc)
        timestamp = now.strftime('%Y-%m-%d__%I:%M:%p').lower().replace(' ', '')
        filename = f'application_contacts__{timestamp}.docx'

        # Save to BytesIO
        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        # Return as HTTP response
        response = HttpResponse(
            file_stream.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'

        return response

    except Exception:
        return generic_500(request=request)
