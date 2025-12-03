from docx import Document
from docx.shared import Pt, RGBColor

from core.models.application import Application
from core.models.procedure_step import ProcedureStepType
from core.views.application.utilities.application_docx_helpers import format_list_with_limit


def add_groups_and_dependencies_section(document: Document, application: Application) -> None:
    """Add application groups and dependencies sections."""
    # Application Groups
    if application.application_groups.exists():
        document.add_heading('Application Groups', level=2)
        format_list_with_limit(document, application.application_groups.all(), limit=20)

    # Dependencies
    has_dependencies = (
        application.application_upstream_dependencies.exists()
        or application.application_downstream_dependencies.exists()
    )

    if has_dependencies:
        document.add_heading('Dependencies', level=2)

        if application.application_upstream_dependencies.exists():
            document.add_heading('Upstream Dependencies', level=3)
            format_list_with_limit(document, application.application_upstream_dependencies.all(), limit=20)

        if application.application_downstream_dependencies.exists():
            document.add_heading('Downstream Dependencies', level=3)
            format_list_with_limit(document, application.application_downstream_dependencies.all(), limit=20)


def add_tools_and_services_section(document: Document, application: Application) -> None:
    """Add tools and service providers sections."""
    has_tools_or_services = application.tools.exists() or application.service_providers.exists()

    if has_tools_or_services:
        document.add_heading('Tools & Services', level=2)

        if application.tools.exists():
            document.add_heading('Tools', level=3)
            format_list_with_limit(document, application.tools.all(), limit=20)

        if application.service_providers.exists():
            document.add_heading('Service Providers', level=3)
            format_list_with_limit(document, application.service_providers.all(), limit=20)


def add_description_section(document: Document, application: Application) -> None:
    """Add description/comment section."""
    if application.comment:
        document.add_heading('Description', level=2)
        para = document.add_paragraph(application.comment)
        for run in para.runs:
            run.font.size = Pt(9)


def add_procedures_section(document: Document, application: Application) -> None:
    """Add procedures section with all procedure steps."""
    procedure_steps = application.procedure_steps.all()

    if not procedure_steps:
        return

    document.add_heading('Procedures', level=2)

    for step in procedure_steps:
        # Step heading
        step_heading = f"Step {step.order}: {step.name}"
        document.add_heading(step_heading, level=3)

        # Step comment/description
        if step.comment:
            para = document.add_paragraph(step.comment)
            for run in para.runs:
                run.font.size = Pt(9)
            para.add_run()  # Add spacing

        # Step content based on type
        if step.step_type == ProcedureStepType.MARKDOWN:
            content = step.step_data.get('content', '')
            if content:
                para = document.add_paragraph(content)
                for run in para.runs:
                    run.font.size = Pt(9)

        elif step.step_type == ProcedureStepType.CODE:
            code = step.step_data.get('code', '')
            language = step.step_data.get('language', 'code')
            if code:
                # Language label
                lang_para = document.add_paragraph()
                lang_run = lang_para.add_run(f'Language: {language}')
                lang_run.font.size = Pt(8)
                lang_run.italic = True
                lang_run.font.color.rgb = RGBColor(100, 100, 100)

                # Code block
                code_para = document.add_paragraph(code)
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(51, 51, 51)

        elif step.step_type == ProcedureStepType.FILE_REFERENCE:
            file_path = step.step_data.get('file_path', '')
            if file_path:
                para = document.add_paragraph()
                run = para.add_run(f'File: {file_path}')
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 51, 51)

        elif step.step_type == ProcedureStepType.CHECKLIST:
            items = step.step_data.get('items', [])
            if items:
                for item in items:
                    item_text = item.get('text', '')
                    if item_text:
                        para = document.add_paragraph(style='List Bullet')
                        run = para.add_run(item_text)
                        run.font.size = Pt(9)
