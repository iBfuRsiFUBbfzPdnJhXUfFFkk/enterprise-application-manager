from docx import Document
from docx.shared import Pt, RGBColor

from core.models.application import Application
from core.models.procedure_step import ProcedureStepType
from core.views.application.utilities.application_docx_helpers import (
    format_list_with_limit,
    format_table_cell,
)


def add_basic_info_section(document: Document, application: Application) -> None:
    """Add basic information 2-column table."""
    document.add_heading('Basic Information', level=2)

    table = document.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'

    format_table_cell(table.rows[0].cells[0], 'Name', bold=True)
    format_table_cell(table.rows[0].cells[1], application.name)

    format_table_cell(table.rows[1].cells[0], 'Acronym', bold=True)
    format_table_cell(table.rows[1].cells[1], application.acronym)

    format_table_cell(table.rows[2].cells[0], 'Aliases', bold=True)
    format_table_cell(table.rows[2].cells[1], application.aliases_csv)

    format_table_cell(table.rows[3].cells[0], 'Lifecycle', bold=True)
    format_table_cell(table.rows[3].cells[1], application.type_lifecycle)

    format_table_cell(table.rows[4].cells[0], 'Launch Date', bold=True)
    format_table_cell(table.rows[4].cells[1], application.date_launch)

    format_table_cell(table.rows[5].cells[0], 'Platform Group', bold=True)
    format_table_cell(table.rows[5].cells[1], application.application_group_platform)


def add_technical_config_section(document: Document, application: Application) -> None:
    """Add technical configuration 2-column table."""
    document.add_heading('Technical Configuration', level=2)

    table = document.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'

    format_table_cell(table.rows[0].cells[0], 'Platform Target', bold=True)
    format_table_cell(table.rows[0].cells[1], application.type_platform_target)

    format_table_cell(table.rows[1].cells[0], 'Deployment Medium', bold=True)
    format_table_cell(table.rows[1].cells[1], application.type_deployment_medium)

    format_table_cell(table.rows[2].cells[0], 'Authentication', bold=True)
    format_table_cell(table.rows[2].cells[1], application.type_authentication)

    format_table_cell(table.rows[3].cells[0], 'Authorization', bold=True)
    format_table_cell(table.rows[3].cells[1], application.type_authorization)

    format_table_cell(table.rows[4].cells[0], 'Peak Userbase', bold=True)
    format_table_cell(table.rows[4].cells[1], application.peak_userbase)

    format_table_cell(table.rows[5].cells[0], 'Pipeline Status', bold=True)
    format_table_cell(table.rows[5].cells[1], application.pipeline_status)

    format_table_cell(table.rows[6].cells[0], 'Centralized Logging', bold=True)
    format_table_cell(table.rows[6].cells[1], application.centralized_logging_status)


def add_compliance_section(document: Document, application: Application) -> None:
    """Add features and compliance bulleted list (only items that are True)."""
    compliance_items = []

    if application.is_externally_facing:
        compliance_items.append('Externally Facing')
    if application.is_legacy:
        compliance_items.append('Legacy Application')
    if application.is_using_artificial_intelligence:
        compliance_items.append('Uses Artificial Intelligence')
    if application.is_storing_personally_identifiable_information_pii:
        compliance_items.append('Stores PII (Personally Identifiable Information)')
    if application.is_storing_protected_health_information_phi:
        compliance_items.append('Stores PHI (Protected Health Information)')
    if application.is_storing_nonpublic_personal_information_npi:
        compliance_items.append('Stores NPI (Nonpublic Personal Information)')
    if application.is_required_to_adhere_to_general_data_protection_regulation_gdpr:
        compliance_items.append('GDPR Compliance Required')
    if application.is_required_to_adhere_to_california_consumer_privacy_act_ccpa:
        compliance_items.append('CCPA Compliance Required')
    if application.is_required_to_adhere_to_payment_card_industry_data_security_standard_pci_dss:
        compliance_items.append('PCI-DSS Compliance Required')

    if compliance_items:
        document.add_heading('Features & Compliance', level=2)
        for item in compliance_items:
            para = document.add_paragraph(style='List Bullet')
            run = para.add_run(item)
            run.font.size = Pt(9)


def add_team_section(document: Document, application: Application) -> None:
    """Add team section with key roles, developers, SMEs, and stakeholders."""
    # Check if any team members exist
    has_team = any([
        application.person_product_owner,
        application.person_product_manager,
        application.person_project_manager,
        application.person_scrum_master,
        application.person_architect,
        application.person_lead_developer,
        application.person_developers.exists(),
        application.person_smes.exists(),
        application.person_stakeholders.exists(),
    ])

    if not has_team:
        return

    document.add_heading('Team', level=2)

    # Key roles table
    table = document.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'

    format_table_cell(table.rows[0].cells[0], 'Product Owner', bold=True)
    format_table_cell(table.rows[0].cells[1], application.person_product_owner)

    format_table_cell(table.rows[1].cells[0], 'Product Manager', bold=True)
    format_table_cell(table.rows[1].cells[1], application.person_product_manager)

    format_table_cell(table.rows[2].cells[0], 'Project Manager', bold=True)
    format_table_cell(table.rows[2].cells[1], application.person_project_manager)

    format_table_cell(table.rows[3].cells[0], 'Scrum Master', bold=True)
    format_table_cell(table.rows[3].cells[1], application.person_scrum_master)

    format_table_cell(table.rows[4].cells[0], 'Architect', bold=True)
    format_table_cell(table.rows[4].cells[1], application.person_architect)

    format_table_cell(table.rows[5].cells[0], 'Lead Developer', bold=True)
    format_table_cell(table.rows[5].cells[1], application.person_lead_developer)

    # Developers
    if application.person_developers.exists():
        document.add_heading('Developers', level=3)
        format_list_with_limit(document, application.person_developers.all(), limit=20)

    # SMEs
    if application.person_smes.exists():
        document.add_heading('Subject Matter Experts (SMEs)', level=3)
        format_list_with_limit(document, application.person_smes.all(), limit=20)

    # Stakeholders
    if application.person_stakeholders.exists():
        document.add_heading('Stakeholders', level=3)
        format_list_with_limit(document, application.person_stakeholders.all(), limit=20)


def add_links_section(document: Document, application: Application) -> None:
    """Add links table with all populated link fields."""
    link_fields = [
        ('Production Server', application.link_production_server),
        ('Production Server (External)', application.link_production_server_external),
        ('Staging Server', application.link_staging_server),
        ('Development Server', application.link_development_server),
        ('GitLab Repository', application.link_gitlab_repository),
        ('GitLab Wiki', application.link_gitlab_wiki),
        ('Teams Channel', application.link_teams_channel),
        ('Logs', application.link_logs),
        ('Sentry.io', application.link_sentry_io),
        ('Postman', application.link_postman),
        ('OpenAI', application.link_open_ai),
        ('Divio', application.link_divio),
        ('Lucid', application.link_lucid),
        ('Whiteboard', application.link_whiteboard),
        ('Wrike', application.link_wrike),
        ('Support Email', application.link_support_email),
        ('Training', application.link_training),
        ('Ticket Submission', application.link_ticket_submission),
        ('SBOM', application.link_software_bill_of_materials_sbom),
    ]

    populated_links = [(label, url) for label, url in link_fields if url]

    if populated_links:
        document.add_heading('Links', level=2)
        table = document.add_table(rows=len(populated_links), cols=2)
        table.style = 'Light Grid Accent 1'

        for i, (label, url) in enumerate(populated_links):
            format_table_cell(table.rows[i].cells[0], label, bold=True)
            format_table_cell(table.rows[i].cells[1], url)
