from datetime import datetime, timezone
from io import BytesIO
from collections import OrderedDict
from decimal import Decimal

from django.http import HttpRequest, HttpResponse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.models.estimation import Estimation
from core.views.generic.generic_500 import generic_500


def estimation_export_docx_view(request: HttpRequest, model_id: int) -> HttpResponse:
    """
    Export estimation to DOCX document.
    """
    try:
        estimation = Estimation.objects.get(id=model_id)
    except Estimation.DoesNotExist:
        return generic_500(request=request)

    # Create a new Document
    document = Document()

    # Set narrow margins (0.5 inches on all sides)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Add title
    title = document.add_heading(estimation.name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add description if present
    if estimation.description:
        document.add_paragraph(estimation.description)
        document.add_paragraph()

    # Add methodology explanation section
    document.add_heading('Estimation Methodology', level=2)

    # Overview paragraph
    overview = document.add_paragraph()
    overview.add_run('This estimation uses a sophisticated approach based on the Cone of Uncertainty principle to provide realistic project timelines. ').font.size = Pt(10)
    overview.add_run('The methodology applies uncertainty multipliers based on project phase and adds contingency padding for unforeseen challenges.').font.size = Pt(10)
    document.add_paragraph()

    # Cone of Uncertainty explanation
    document.add_heading('Cone of Uncertainty', level=3)
    cou_para = document.add_paragraph()
    cou_para.add_run('The Cone of Uncertainty recognizes that estimates become more accurate as a project progresses through its lifecycle. ').font.size = Pt(10)
    cou_para.add_run('Each task is assigned an uncertainty multiplier based on its current phase:').font.size = Pt(10)

    cou_list = [
        ('Initial Concept (4x):', 'Early exploration phase with highest uncertainty. Requirements are vague and architecture is undefined.'),
        ('Approved Product (2x):', 'Product definition complete, but detailed requirements are still being refined and technical approach is being validated.'),
        ('Requirements Complete (1.5x):', 'Requirements are documented and approved, design work is in progress, but implementation details may vary.'),
        ('Design Complete (1.25x):', 'Technical design is finished and implementation is starting. Most unknowns have been resolved.'),
        ('Implementation Complete (1.1x):', 'Code is complete and being tested. Only minor refinements and bug fixes remain.'),
    ]

    for label, description in cou_list:
        p = document.add_paragraph(style='List Bullet')
        run = p.add_run(f'{label} ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(description)
        run2.font.size = Pt(10)

    document.add_paragraph()

    # Separated hours components
    document.add_heading('Hour Components', level=3)
    components_para = document.add_paragraph()
    components_para.add_run('Each estimation item separates work into distinct components:').font.size = Pt(10)

    components_list = [
        ('Development Hours:', 'Core implementation time for writing features, fixing bugs, or building functionality.'),
        ('Code Review Hours:', 'Time spent reviewing others\' code (typically 0.5x development time). Critical for maintaining code quality and catching issues early.'),
        ('Testing Hours:', 'Time for writing unit tests, integration tests, and manual QA (typically 1x development time). Essential for long-term maintainability.'),
        ('Code Reviewer Hours:', 'Time spent in the back-and-forth code review process - addressing feedback, resolving issues, and iterating on changes based on review comments.'),
    ]

    for label, description in components_list:
        p = document.add_paragraph(style='List Bullet')
        run = p.add_run(f'{label} ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(description)
        run2.font.size = Pt(10)

    document.add_paragraph()

    # Contingency padding
    document.add_heading('Contingency Padding', level=3)
    contingency_para = document.add_paragraph()
    contingency_value = float(estimation.contingency_padding_percent or 0)
    contingency_para.add_run(f'This estimation includes a {contingency_value:.1f}% contingency buffer applied to base hours (before uncertainty multipliers). ').font.size = Pt(10)
    contingency_para.add_run('This padding accounts for:').font.size = Pt(10)

    contingency_list = [
        'Scope creep and requirement changes',
        'Integration challenges with existing systems',
        'Unforeseen technical debt or refactoring needs',
        'Team coordination overhead and communication time',
        'Deployment, documentation, and release activities',
    ]

    for item in contingency_list:
        p = document.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(10)

    document.add_paragraph()

    # Calculation formula
    document.add_heading('Calculation Formula', level=3)
    formula_para = document.add_paragraph()
    formula_para.add_run('The total hours for each task are calculated as follows:').font.size = Pt(10)
    document.add_paragraph()

    formula_steps = [
        '1. Base Hours = Development + Code Review + Testing',
        '2. Hours with Uncertainty = Base Hours × Cone of Uncertainty Multiplier',
        '3. Contingency Hours = Base Hours × Contingency Percentage',
        '4. Total Hours = Hours with Uncertainty + Contingency Hours',
    ]

    for step in formula_steps:
        p = document.add_paragraph()
        run = p.add_run(step)
        run.font.size = Pt(10)
        run.bold = True

    document.add_paragraph()
    example_para = document.add_paragraph()
    example_para.add_run('Example: A mid-level task with 2 dev hours, 1 code review hour, 2 test hours (5 base hours total), Design Complete phase (1.25x), and 20% contingency:').font.size = Pt(10)

    example_steps = [
        '• Base Hours = 5',
        '• With Uncertainty = 5 × 1.25 = 6.25 hours',
        '• Contingency = 5 × 0.20 = 1.0 hour',
        '• Total = 6.25 + 1.0 = 7.25 hours',
    ]

    for step in example_steps:
        p = document.add_paragraph()
        run = p.add_run(step)
        run.font.size = Pt(10)

    document.add_paragraph()

    # How to use estimates
    document.add_heading('Interpreting the Estimates', level=3)
    interpret_para = document.add_paragraph()
    interpret_para.add_run('Each developer level provides hours estimates for that specific experience level. ').font.size = Pt(10)
    interpret_para.add_run('Choose the level that matches your team composition and consider these factors:').font.size = Pt(10)

    interpret_list = [
        'All developer levels use the same cone of uncertainty multipliers—the difference is in base productivity',
        'Junior developers typically require more base hours for the same task compared to senior developers',
        'Contingency padding is applied to base hours to account for unforeseen challenges',
        'Review iteration hours represent time spent addressing feedback and resolving issues from code reviews',
        'Mixed teams should estimate tasks using the appropriate level for who will actually perform the work',
    ]

    for item in interpret_list:
        p = document.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(10)

    document.add_paragraph()

    # Add metadata section
    document.add_heading('Project Information', level=2)
    table = document.add_table(rows=9, cols=2)
    table.style = 'Light Grid Accent 1'

    # Application
    row = table.rows[0]
    row.cells[0].text = 'Application'
    row.cells[1].text = str(estimation.application) if estimation.application else 'N/A'

    # Project
    row = table.rows[1]
    row.cells[0].text = 'Project'
    row.cells[1].text = str(estimation.project) if estimation.project else 'N/A'

    # Contingency Padding
    row = table.rows[2]
    row.cells[0].text = 'Contingency Padding'
    row.cells[1].text = f"{float(estimation.contingency_padding_percent or 0):.2f}%"

    # Total Story Points
    row = table.rows[3]
    row.cells[0].text = 'Total Story Points'
    total_points = estimation.get_total_story_points()
    row.cells[1].text = f"{float(total_points):.1f} points"

    # Sprint Count
    row = table.rows[4]
    row.cells[0].text = 'Sprint Count'
    if estimation.sprint_duration_weeks and estimation.get_total_team_size() > 0:
        sprint_count = estimation.get_sprint_count()
        if sprint_count:
            sprint_text = f"{sprint_count} sprint{'s' if sprint_count != 1 else ''} ({estimation.sprint_duration_weeks} week{'s' if estimation.sprint_duration_weeks != 1 else ''} each)"
            row.cells[1].text = sprint_text
        else:
            row.cells[1].text = 'N/A'
    else:
        if not estimation.sprint_duration_weeks:
            row.cells[1].text = 'N/A (sprint duration not configured)'
        else:
            row.cells[1].text = 'N/A (team composition required)'

    # Required Team Velocity
    row = table.rows[5]
    row.cells[0].text = 'Required Team Velocity'
    team_velocity = estimation.get_required_team_velocity_per_sprint()
    if team_velocity:
        row.cells[1].text = f"{float(team_velocity):.1f} story points per sprint"
    else:
        row.cells[1].text = 'N/A (requires sprint count)'

    # Required Per-Developer Velocity
    row = table.rows[6]
    row.cells[0].text = 'Required Per-Developer Velocity'
    dev_velocity = estimation.get_required_velocity_per_developer_per_sprint()
    if dev_velocity:
        row.cells[1].text = f"{float(dev_velocity):.2f} story points per sprint per developer"
    else:
        row.cells[1].text = 'N/A (requires sprint count and team size)'

    # Required Weekly Velocity
    row = table.rows[7]
    row.cells[0].text = 'Required Weekly Velocity'
    weekly_velocity = estimation.get_required_velocity_per_week()
    if weekly_velocity:
        row.cells[1].text = f"{float(weekly_velocity):.1f} story points per week"
    else:
        row.cells[1].text = 'N/A (requires project duration)'

    # Average Hours
    row = table.rows[8]
    row.cells[0].text = 'Average Hours (All Levels)'
    avg_hours = estimation.get_average_hours_with_uncertainty()
    row.cells[1].text = f"{float(avg_hours):.2f} hours with uncertainty"

    # Make all text in table smaller
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    document.add_paragraph()

    # Add combined project estimate (if team composition is set)
    if estimation.get_total_team_size() > 0:
        document.add_heading('Combined Project Estimate', level=2)

        combined_para = document.add_paragraph()
        combined_para.add_run('Based on your team composition, this is the realistic project timeline assuming all developer levels work in parallel on their assigned tasks. Code review time is included in each developer\'s hours.').font.size = Pt(10)
        document.add_paragraph()

        # Determine number of rows needed (9 rows total)
        num_rows = 9

        # Create table for combined estimate
        combined_table = document.add_table(rows=num_rows, cols=2)
        combined_table.style = 'Light Grid Accent 1'

        row_idx = 0

        # Team size
        row = combined_table.rows[row_idx]
        row.cells[0].text = 'Total Team Size'
        team_text = f"{estimation.get_total_team_size()} developer{'s' if estimation.get_total_team_size() != 1 else ''}"
        row.cells[1].text = team_text
        row_idx += 1

        # Project duration in weeks
        row = combined_table.rows[row_idx]
        row.cells[0].text = 'Project Duration (Weeks)'
        weeks = estimation.get_combined_project_duration_weeks()
        row.cells[1].text = f"{float(weeks):.1f} weeks" if weeks else 'N/A'
        row_idx += 1

        # Project duration in months
        row = combined_table.rows[row_idx]
        row.cells[0].text = 'Project Duration (Months)'
        months = estimation.get_combined_project_duration_months()
        row.cells[1].text = f"{float(months):.1f} months" if months else 'N/A'
        row_idx += 1

        # Sprint count
        row = combined_table.rows[row_idx]
        row.cells[0].text = 'Sprint Count'
        if estimation.sprint_duration_weeks:
            sprint_count = estimation.get_sprint_count()
            if sprint_count:
                sprint_text = f"{sprint_count} sprint{'s' if sprint_count != 1 else ''} ({estimation.sprint_duration_weeks} week{'s' if estimation.sprint_duration_weeks != 1 else ''} each)"
                row.cells[1].text = sprint_text
            else:
                row.cells[1].text = 'N/A'
        else:
            row.cells[1].text = 'N/A (sprint duration not configured)'
        row_idx += 1

        # Required Team Velocity
        row = combined_table.rows[row_idx]
        row.cells[0].text = 'Required Team Velocity'
        team_velocity = estimation.get_required_team_velocity_per_sprint()
        if team_velocity:
            row.cells[1].text = f"{float(team_velocity):.1f} story points per sprint"
        else:
            row.cells[1].text = 'N/A (requires sprint count)'
        row_idx += 1

        # Required Per-Developer Velocity
        row = combined_table.rows[row_idx]
        row.cells[0].text = 'Required Per-Developer Velocity'
        dev_velocity = estimation.get_required_velocity_per_developer_per_sprint()
        if dev_velocity:
            row.cells[1].text = f"{float(dev_velocity):.2f} story points per sprint per developer"
        else:
            row.cells[1].text = 'N/A (requires sprint count and team size)'
        row_idx += 1

        # Required Weekly Velocity
        row = combined_table.rows[row_idx]
        row.cells[0].text = 'Required Weekly Velocity'
        weekly_velocity = estimation.get_required_velocity_per_week()
        if weekly_velocity:
            row.cells[1].text = f"{float(weekly_velocity):.1f} story points per week"
        else:
            row.cells[1].text = 'N/A (requires project duration)'
        row_idx += 1

        # Bottleneck level
        row = combined_table.rows[row_idx]
        row.cells[0].text = 'Bottleneck Level'
        bottleneck = estimation.get_bottleneck_level()
        if bottleneck:
            row.cells[1].text = f"{bottleneck[0]} ({float(bottleneck[1]):.1f} weeks - determines project duration)"
        else:
            row.cells[1].text = 'N/A'
        row_idx += 1

        # Average hours across all levels
        row = combined_table.rows[row_idx]
        row.cells[0].text = 'Average Hours (All Levels)'
        avg_hours = estimation.get_average_hours_with_uncertainty()
        row.cells[1].text = f"{float(avg_hours):.2f} hours with uncertainty"

        # Make all text in table smaller
        for row in combined_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        document.add_paragraph()

        # Team composition breakdown
        document.add_heading('Team Composition', level=3)
        comp_table = document.add_table(rows=4, cols=2)
        comp_table.style = 'Light Grid Accent 1'

        comp_table.rows[0].cells[0].text = 'Junior Developers'
        comp_table.rows[0].cells[1].text = str(estimation.junior_developer_count or 0)

        comp_table.rows[1].cells[0].text = 'Mid-Level Developers'
        comp_table.rows[1].cells[1].text = str(estimation.mid_developer_count or 0)

        comp_table.rows[2].cells[0].text = 'Senior Developers'
        comp_table.rows[2].cells[1].text = str(estimation.senior_developer_count or 0)

        comp_table.rows[3].cells[0].text = 'Lead Developers'
        comp_table.rows[3].cells[1].text = str(estimation.lead_developer_count or 0)

        # Make all text in table smaller
        for row in comp_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        document.add_paragraph()

        # Add Best Case Scenario
        document.add_heading('Best Case Scenario', level=2)

        best_case_para = document.add_paragraph()
        best_case_para.add_run('This optimistic timeline assumes all estimation items are at the "Implementation Complete" phase with minimal uncertainty (1.1x multiplier). ').font.size = Pt(10)
        best_case_para.add_run('This represents the most favorable scenario where requirements, design, and implementation details are fully known.').font.size = Pt(10)
        document.add_paragraph()

        # Create table for best case
        best_case_table = document.add_table(rows=4, cols=2)
        best_case_table.style = 'Light Grid Accent 1'

        row_idx = 0

        # Team size
        row = best_case_table.rows[row_idx]
        row.cells[0].text = 'Total Team Size'
        team_text = f"{estimation.get_total_team_size()} developer{'s' if estimation.get_total_team_size() != 1 else ''}"
        row.cells[1].text = team_text
        row_idx += 1

        # Best case duration in weeks
        row = best_case_table.rows[row_idx]
        row.cells[0].text = 'Project Duration (Weeks)'
        weeks_best = estimation.get_combined_project_duration_weeks_best_case()
        row.cells[1].text = f"{float(weeks_best):.1f} weeks" if weeks_best else 'N/A'
        row_idx += 1

        # Best case duration in months
        row = best_case_table.rows[row_idx]
        row.cells[0].text = 'Project Duration (Months)'
        months_best = estimation.get_combined_project_duration_months_best_case()
        row.cells[1].text = f"{float(months_best):.1f} months" if months_best else 'N/A'
        row_idx += 1

        # Uncertainty multiplier
        row = best_case_table.rows[row_idx]
        row.cells[0].text = 'Uncertainty Multiplier'
        row.cells[1].text = '1.1x (Implementation Complete)'

        # Make all text in table smaller
        for row in best_case_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        document.add_paragraph()

        # Add Worst Case Scenario
        document.add_heading('Worst Case Scenario', level=2)

        worst_case_para = document.add_paragraph()
        worst_case_para.add_run('This conservative timeline assumes all estimation items are at the "Initial Concept" stage with maximum uncertainty (4.0x multiplier). ').font.size = Pt(10)
        worst_case_para.add_run('This represents the most conservative scenario where requirements are unclear and significant changes are expected during the project.').font.size = Pt(10)
        document.add_paragraph()

        # Create table for worst case
        worst_case_table = document.add_table(rows=4, cols=2)
        worst_case_table.style = 'Light Grid Accent 1'

        row_idx = 0

        # Team size
        row = worst_case_table.rows[row_idx]
        row.cells[0].text = 'Total Team Size'
        team_text = f"{estimation.get_total_team_size()} developer{'s' if estimation.get_total_team_size() != 1 else ''}"
        row.cells[1].text = team_text
        row_idx += 1

        # Worst case duration in weeks
        row = worst_case_table.rows[row_idx]
        row.cells[0].text = 'Project Duration (Weeks)'
        weeks_worst = estimation.get_combined_project_duration_weeks_worst_case()
        row.cells[1].text = f"{float(weeks_worst):.1f} weeks" if weeks_worst else 'N/A'
        row_idx += 1

        # Worst case duration in months
        row = worst_case_table.rows[row_idx]
        row.cells[0].text = 'Project Duration (Months)'
        months_worst = estimation.get_combined_project_duration_months_worst_case()
        row.cells[1].text = f"{float(months_worst):.1f} months" if months_worst else 'N/A'
        row_idx += 1

        # Uncertainty multiplier
        row = worst_case_table.rows[row_idx]
        row.cells[0].text = 'Uncertainty Multiplier'
        row.cells[1].text = '4.0x (Initial Concept)'

        # Make all text in table smaller
        for row in worst_case_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        document.add_paragraph()

    # Add estimation items section
    document.add_heading('Estimation Items', level=2)
    document.add_paragraph('Hours shown include cone of uncertainty multipliers based on project phase. Each item\'s base hours (dev + code review + testing) are multiplied by its cone of uncertainty factor. Code review is performed by the same developers, not a separate reviewer position.')

    items = estimation.items.all().order_by('order', 'id')

    # Group items by their group field with subtotals
    grouped_items = OrderedDict()
    grouped_items['Ungrouped'] = {'items': [], 'subtotals': {}}

    for item in items:
        group_name = item.group if item.group else 'Ungrouped'
        if group_name not in grouped_items:
            grouped_items[group_name] = {'items': [], 'subtotals': {}}
        grouped_items[group_name]['items'].append(item)

    # Remove Ungrouped if it's empty
    if not grouped_items['Ungrouped']['items']:
        del grouped_items['Ungrouped']

    # Calculate subtotals for each group
    for group_name, group_data in grouped_items.items():
        group_items_list = group_data['items']
        subtotals = {
            'story_points': sum(item.story_points or Decimal('0') for item in group_items_list),
            'junior_with_uncertainty': sum(item.get_junior_hours_with_uncertainty() for item in group_items_list),
            'mid_with_uncertainty': sum(item.get_mid_hours_with_uncertainty() for item in group_items_list),
            'senior_with_uncertainty': sum(item.get_senior_hours_with_uncertainty() for item in group_items_list),
            'lead_with_uncertainty': sum(item.get_lead_hours_with_uncertainty() for item in group_items_list),
            'reviewer_hours': sum(item.get_reviewer_hours() for item in group_items_list),
        }
        group_data['subtotals'] = subtotals

    if items:
        # Create table for items (hours with uncertainty applied)
        table = document.add_table(rows=1, cols=10)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'ID'
        header_cells[1].text = 'Title'
        header_cells[2].text = 'Pts'
        header_cells[3].text = 'Cmplx'
        header_cells[4].text = 'Prior'
        header_cells[5].text = 'CoU'
        header_cells[6].text = 'Jr'
        header_cells[7].text = 'Mid'
        header_cells[8].text = 'Sr'
        header_cells[9].text = 'Lead'

        # Make header bold and smaller font
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Add items grouped with headers and subtotals
        item_counter = 1
        for group_name, group_data in grouped_items.items():
            # Add group header row
            group_row_cells = table.add_row().cells
            # Merge all cells for group header
            group_row_cells[0].text = f"📁 {group_name}"
            group_row_cells[0].merge(group_row_cells[9])

            # Make group header bold and slightly larger
            for paragraph in group_row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

            # Add items in this group
            for item in group_data['items']:
                row_cells = table.add_row().cells

                # Set text content with ID as first column
                item_ref = f"ITEM-{item_counter:03d}"
                contents = [
                    item_ref,
                    item.title if item.title else '(No title)',
                    f"{float(item.story_points or 0):.1f}",
                    item.get_complexity_level_display() if item.complexity_level else 'N/A',
                    item.get_priority_display() if item.priority else 'N/A',
                    item.get_cone_of_uncertainty_display() if item.cone_of_uncertainty else 'N/A',
                    f"{float(item.get_junior_hours_with_uncertainty()):.2f}",
                    f"{float(item.get_mid_hours_with_uncertainty()):.2f}",
                    f"{float(item.get_senior_hours_with_uncertainty()):.2f}",
                    f"{float(item.get_lead_hours_with_uncertainty()):.2f}"
                ]

                for col_idx, content in enumerate(contents):
                    row_cells[col_idx].text = content
                    # Set font size for the paragraph
                    for paragraph in row_cells[col_idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

                item_counter += 1

            # Add subtotal row for this group
            subtotal_row_cells = table.add_row().cells
            subtotal_row_cells[0].text = f"{group_name} Subtotal"
            subtotal_row_cells[1].text = ""
            subtotal_row_cells[2].text = f"{float(group_data['subtotals']['story_points']):.1f}"
            subtotal_row_cells[3].text = ""
            subtotal_row_cells[4].text = ""
            subtotal_row_cells[5].text = ""
            subtotal_row_cells[6].text = f"{float(group_data['subtotals']['junior_with_uncertainty']):.2f}"
            subtotal_row_cells[7].text = f"{float(group_data['subtotals']['mid_with_uncertainty']):.2f}"
            subtotal_row_cells[8].text = f"{float(group_data['subtotals']['senior_with_uncertainty']):.2f}"
            subtotal_row_cells[9].text = f"{float(group_data['subtotals']['lead_with_uncertainty']):.2f}"

            # Make subtotal row bold
            for cell in subtotal_row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)

    else:
        document.add_paragraph('No estimation items added yet.')

    document.add_paragraph()

    # Add summary section
    document.add_heading('Estimation Summary', level=2)
    document.add_paragraph('Hours include cone of uncertainty multipliers based on project phase. Contingency is applied to base hours only. Each level represents alternative estimates, not additive totals. Code review time is included in each developer\'s hours.')

    summary_table = document.add_table(rows=19, cols=2)
    summary_table.style = 'Light Grid Accent 1'

    # Calculate totals per level
    totals = [
        ('JUNIOR DEVELOPER', ''),
        ('  Base Hours', estimation.get_base_hours_junior()),
        ('  With Uncertainty', estimation.get_total_hours_junior_with_uncertainty()),
        (f'  Contingency ({float(estimation.contingency_padding_percent or 0):.2f}%)', estimation.get_contingency_hours_junior()),
        ('  Grand Total', estimation.get_grand_total_hours_junior()),
        ('MID-LEVEL DEVELOPER', ''),
        ('  Base Hours', estimation.get_base_hours_mid()),
        ('  With Uncertainty', estimation.get_total_hours_mid_with_uncertainty()),
        (f'  Contingency ({float(estimation.contingency_padding_percent or 0):.2f}%)', estimation.get_contingency_hours_mid()),
        ('  Grand Total', estimation.get_grand_total_hours_mid()),
        ('SENIOR DEVELOPER', ''),
        ('  Base Hours', estimation.get_base_hours_senior()),
        ('  With Uncertainty', estimation.get_total_hours_senior_with_uncertainty()),
        (f'  Contingency ({float(estimation.contingency_padding_percent or 0):.2f}%)', estimation.get_contingency_hours_senior()),
        ('  Grand Total', estimation.get_grand_total_hours_senior()),
        ('LEAD DEVELOPER', ''),
        ('  Base Hours', estimation.get_base_hours_lead()),
        ('  With Uncertainty', estimation.get_total_hours_lead_with_uncertainty()),
        (f'  Contingency ({float(estimation.contingency_padding_percent or 0):.2f}%)', estimation.get_contingency_hours_lead()),
        ('  Grand Total', estimation.get_grand_total_hours_lead()),
        ('', ''),  # Empty row for spacing
        ('TOTAL STORY POINTS', estimation.get_total_story_points()),
    ]

    # Adjust rows to match the number of totals
    while len(summary_table.rows) < len(totals):
        summary_table.add_row()

    for idx, (label, value) in enumerate(totals):
        row = summary_table.rows[idx]
        row.cells[0].text = label
        if value == '':
            # Section headers or empty rows
            row.cells[1].text = ''
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
        else:
            # Format based on whether it's story points or hours
            if 'STORY POINTS' in label:
                row.cells[1].text = f"{float(value):.1f} points"
            else:
                row.cells[1].text = f"{float(value):.2f} hours"

            # Make grand total and story points rows bold
            if 'Grand Total' in label or 'STORY POINTS' in label:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)
            else:
                # Regular rows - make smaller
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

    # Add page break and detailed items section
    if items:
        document.add_page_break()
        document.add_heading('Detailed Item Descriptions', level=1)
        detail_para = document.add_paragraph()
        detail_para.add_run('This section provides detailed information for each estimation item, including full descriptions. ').font.size = Pt(10)
        detail_para.add_run('Items are organized by group. Reference the item ID to cross-reference with the summary table above.').font.size = Pt(10)
        document.add_paragraph()

        # Add detailed entry for each item grouped by category
        item_counter = 1
        for group_name, group_data in grouped_items.items():
            # Add group header
            group_heading = document.add_heading(f"📁 {group_name}", level=2)
            document.add_paragraph()

            # Add items in this group
            for item in group_data['items']:
                item_ref = f"ITEM-{item_counter:03d}"

                # Add item reference ID as heading
                item_heading = document.add_heading(f"{item_ref}: {item.title if item.title else '(No title)'}", level=3)

                # Add metadata table for the item
                meta_table = document.add_table(rows=6, cols=2)
                meta_table.style = 'Light Grid Accent 1'

                meta_table.rows[0].cells[0].text = 'Story Points'
                meta_table.rows[0].cells[1].text = f"{float(item.story_points or 0):.1f}"

                meta_table.rows[1].cells[0].text = 'Complexity'
                meta_table.rows[1].cells[1].text = item.get_complexity_level_display() if item.complexity_level else 'N/A'

                meta_table.rows[2].cells[0].text = 'Priority'
                meta_table.rows[2].cells[1].text = item.get_priority_display() if item.priority else 'N/A'

                meta_table.rows[3].cells[0].text = 'Cone of Uncertainty'
                meta_table.rows[3].cells[1].text = item.get_cone_of_uncertainty_display() if item.cone_of_uncertainty else 'N/A'

                # Add hours breakdown
                meta_table.rows[4].cells[0].text = 'Hours (Jr/Mid/Sr/Lead)'
                hours_str = f"{float(item.get_junior_hours_with_uncertainty()):.1f} / {float(item.get_mid_hours_with_uncertainty()):.1f} / {float(item.get_senior_hours_with_uncertainty()):.1f} / {float(item.get_lead_hours_with_uncertainty()):.1f}"
                meta_table.rows[4].cells[1].text = hours_str

                # Add base hours
                meta_table.rows[5].cells[0].text = 'Base Hours (Lead)'
                base_hours = (item.hours_lead or 0) + (item.code_review_hours_lead or 0) + (item.tests_hours_lead or 0)
                meta_table.rows[5].cells[1].text = f"{float(base_hours):.1f}"

                # Make all text in table smaller
                for row in meta_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)

                document.add_paragraph()

                # Add description section
                if item.description:
                    document.add_heading('Description', level=4)
                    # Basic markdown rendering
                    description_lines = item.description.split('\n')

                    for line in description_lines:
                        line = line.strip()
                        if not line:
                            document.add_paragraph()
                            continue

                        # Handle headings
                        if line.startswith('# '):
                            document.add_heading(line[2:], level=5)
                        elif line.startswith('## '):
                            document.add_heading(line[3:], level=6)
                        elif line.startswith('### '):
                            document.add_heading(line[4:], level=6)
                        # Handle bullet lists
                        elif line.startswith('- ') or line.startswith('* '):
                            p = document.add_paragraph(line[2:], style='List Bullet')
                            for run in p.runs:
                                run.font.size = Pt(10)
                        # Handle numbered lists
                        elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
                            p = document.add_paragraph(line[3:], style='List Number')
                            for run in p.runs:
                                run.font.size = Pt(10)
                        # Handle code blocks
                        elif line.startswith('```'):
                            continue  # Skip code fence markers
                        # Regular paragraph
                        else:
                            p = document.add_paragraph()
                            # Simple inline markdown: **bold** and *italic*
                            import re
                            # Replace **bold**
                            parts = re.split(r'(\*\*.*?\*\*)', line)
                            for part in parts:
                                if part.startswith('**') and part.endswith('**'):
                                    run = p.add_run(part[2:-2])
                                    run.bold = True
                                    run.font.size = Pt(10)
                                else:
                                    # Replace *italic*
                                    italic_parts = re.split(r'(\*.*?\*)', part)
                                    for italic_part in italic_parts:
                                        if italic_part.startswith('*') and italic_part.endswith('*') and not italic_part.startswith('**'):
                                            run = p.add_run(italic_part[1:-1])
                                            run.italic = True
                                            run.font.size = Pt(10)
                                        else:
                                            run = p.add_run(italic_part)
                                            run.font.size = Pt(10)
                else:
                    document.add_paragraph('No description provided.')

                document.add_paragraph()  # Add spacing between items
                item_counter += 1

    # Save document to BytesIO
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Generate filename: <name>__<YYYY-MM-DD>__<HH:MM:[AM|PM]>.docx (lowercase, spaces to underscores, UTC time, 12-hour format)
    now = datetime.now(timezone.utc)
    date_str = now.strftime('%Y-%m-%d')
    time_str = now.strftime('%I:%M:%p').lower()  # 12-hour format with lowercase am/pm
    name_cleaned = estimation.name.lower().replace(' ', '_')
    filename = f"{name_cleaned}__{date_str}__{time_str}.docx"

    # Create HTTP response
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    return response
